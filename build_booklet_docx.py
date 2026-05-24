#!/usr/bin/env python3
"""
STEP 2b — Build Booklet Word Document (two-pass, native footnotes)

Pass 1: Build full .docx with superscript number placeholders.
Pass 2: Re-open zip, swap every superscript-number run for a real
        <w:footnoteReference>, inject word/footnotes.xml, patch rels.
"""

import json, re, argparse, zipfile, shutil, os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from lxml import etree

PAGE_SIZES = {
    'half-letter': (5.5, 8.5),
    'a5':          (5.83, 8.27),
    'letter':      (8.5, 11.0),
}
MARGINS = {
    'half-letter': {'top': 0.75, 'bottom': 0.75, 'inner': 0.875, 'outer': 0.625, 'header': 0.4, 'footer': 0.4},
    'a5':          {'top': 0.75, 'bottom': 0.75, 'inner': 0.875, 'outer': 0.625, 'header': 0.4, 'footer': 0.4},
    'letter':      {'top': 1.0,  'bottom': 1.0,  'inner': 1.25,  'outer': 1.0,   'header': 0.5, 'footer': 0.5},
}

_WNS       = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_RNS       = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
_RELS_NS   = 'http://schemas.openxmlformats.org/package/2006/relationships'
_CT_NS     = 'http://schemas.openxmlformats.org/package/2006/content-types'
_XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

_RT_FOOTNOTES = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
_CT_FOOTNOTES = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'


def _w(tag):
    return f'{{{_WNS}}}{tag}'


# ---------------------------------------------------------------------------
# Low-level XML / page-setup helpers
# ---------------------------------------------------------------------------

def set_page_size(section, w_in, h_in):
    sectPr = section._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz'); sectPr.append(pgSz)
    pgSz.set(qn('w:w'), str(int(w_in * 1440)))
    pgSz.set(qn('w:h'), str(int(h_in * 1440)))


def set_mirror_margins(section, m):
    sectPr = section._sectPr
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is None:
        pgMar = OxmlElement('w:pgMar'); sectPr.append(pgMar)
    pgMar.set(qn('w:top'),    str(int(m['top']    * 1440)))
    pgMar.set(qn('w:bottom'), str(int(m['bottom'] * 1440)))
    pgMar.set(qn('w:left'),   str(int(m['inner']  * 1440)))
    pgMar.set(qn('w:right'),  str(int(m['outer']  * 1440)))
    pgMar.set(qn('w:gutter'), '0')
    pgMar.set(qn('w:header'), str(int(m['header'] * 1440)))
    pgMar.set(qn('w:footer'), str(int(m['footer'] * 1440)))


def enable_mirror_margins(doc):
    settings = doc.settings.element
    if settings.find(qn('w:mirrorMargins')) is None:
        settings.append(OxmlElement('w:mirrorMargins'))


def add_page_number_to_footer(footer, align=WD_ALIGN_PARAGRAPH.CENTER):
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = align
    run = para.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    ins = OxmlElement('w:instrText')
    ins.set(qn('xml:space'), 'preserve'); ins.text = ' PAGE '
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(ins); run._r.append(f2)


def set_header_text(header, left_text='', right_text='', center_text=''):
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bot); pPr.append(pBdr)
    if left_text and right_text:
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right'); tab.set(qn('w:pos'), '9000')
        tabs.append(tab); pPr.append(tabs)
        for text, is_tab in [(left_text, False), ('\t', True), (right_text, False)]:
            r = para.add_run(text)
            r.font.name = 'Times New Roman'; r.font.size = Pt(8)
            if not is_tab:
                r.font.italic = True; r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    elif center_text:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(center_text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(8)
        r.font.italic = True; r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def style_run(run, bold=False, italic=False, size_pt=None, font='Times New Roman'):
    run.font.name = font; run.bold = bold; run.italic = italic
    if size_pt: run.font.size = Pt(size_pt)


def para_spacing(para, before_pt=0, after_pt=6, line_rule=None, line_val=None):
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None: sp = OxmlElement('w:spacing'); pPr.append(sp)
    if before_pt is not None: sp.set(qn('w:before'), str(int(before_pt * 20)))
    if after_pt  is not None: sp.set(qn('w:after'),  str(int(after_pt  * 20)))
    if line_rule and line_val:
        sp.set(qn('w:lineRule'), line_rule); sp.set(qn('w:line'), str(line_val))


def para_keep_with_next(para):
    para._p.get_or_add_pPr().append(OxmlElement('w:keepNext'))


# ---------------------------------------------------------------------------
# Inline markdown → docx runs  (Pass 1 placeholders)
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r'(\*\*\*(.+?)\*\*\*'
    r'|\*\*(.+?)\*\*'
    r'|\*(.+?)\*'
    r'|\^(\d{1,3})'
    r'|(?<=[.,;:!?\'\"\)])\s*(\d{1,3})(?=\s|$)'
    r')'
)


def add_inline_runs(para, text, base_size_pt=11, base_font='Times New Roman'):
    """Build runs for para. Footnote refs → superscript placeholder for Pass 2."""
    last = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > last:
            style_run(para.add_run(text[last:m.start()]),
                      size_pt=base_size_pt, font=base_font)
        bi, b, i, fn, fn2 = m.group(2), m.group(3), m.group(4), m.group(5), m.group(6)
        fn = fn or fn2
        if bi:
            style_run(para.add_run(bi), bold=True, italic=True,
                      size_pt=base_size_pt, font=base_font)
        elif b:
            style_run(para.add_run(b), bold=True,
                      size_pt=base_size_pt, font=base_font)
        elif i:
            style_run(para.add_run(i), italic=True,
                      size_pt=base_size_pt, font=base_font)
        elif fn:
            run = para.add_run(fn)
            run.font.name = base_font; run.font.size = Pt(7)
            rPr = run._r.get_or_add_rPr()
            va = OxmlElement('w:vertAlign'); va.set(qn('w:val'), 'superscript')
            rPr.append(va)
        last = m.end()
    if last < len(text):
        style_run(para.add_run(text[last:]), size_pt=base_size_pt, font=base_font)


# ---------------------------------------------------------------------------
# Pass 1 document builders
# ---------------------------------------------------------------------------

def build_title_page(doc, meta, size_name):
    is_small = size_name != 'letter'
    h1 = 16 if is_small else 18
    doc.add_paragraph()
    for _ in range(4): doc.add_paragraph()
    if meta.get('doc_type'):
        p = doc.add_paragraph(meta['doc_type'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = 'Times New Roman'; p.runs[0].font.size = Pt(10)
        para_spacing(p, before_pt=0, after_pt=4)
    title = meta.get('title') or meta.get('latin_title', '')
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.name = 'Times New Roman'; r.font.size = Pt(h1)
        r.bold = True; r.italic = True
        para_spacing(p, before_pt=6, after_pt=10)
    if meta.get('author'):
        p = doc.add_paragraph(meta['author'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(11); p.runs[0].bold = True
        para_spacing(p, before_pt=4, after_pt=4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr'); bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1');   bot.set(qn('w:color'), '888888')
    pBdr.append(bot); pPr.append(pBdr)
    para_spacing(p, before_pt=8, after_pt=8)
    for line in meta.get('addressees', '').split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.name = 'Times New Roman'; p.runs[0].font.size = Pt(10)
            para_spacing(p, before_pt=0, after_pt=3)
    if meta.get('subject'):
        p = doc.add_paragraph(meta['subject'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(11); p.runs[0].italic = True
        para_spacing(p, before_pt=10, after_pt=4)
    doc.add_page_break()


def build_body(doc, chapters, size_name):
    is_small = size_name != 'letter'
    body_pt = 10.5 if is_small else 11
    ch_pt   = 13   if is_small else 14
    for chapter in chapters:
        title    = chapter.get('title', '')
        subtitle = chapter.get('subtitle', '')
        number   = chapter.get('number', 0)
        paras    = chapter.get('paragraphs', [])
        if not paras and not title:
            continue
        if title:
            if number and number > 0:
                p = doc.add_paragraph(f'CHAPTER {number}')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.name = 'Times New Roman'; p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                para_spacing(p, before_pt=18, after_pt=2); para_keep_with_next(p)
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p2.add_run(subtitle or title)
                r.font.name = 'Times New Roman'; r.font.size = Pt(ch_pt); r.bold = True
                para_spacing(p2, before_pt=0, after_pt=14); para_keep_with_next(p2)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(title.upper())
                r.font.name = 'Times New Roman'; r.font.size = Pt(ch_pt - 1); r.bold = True
                para_spacing(p, before_pt=14, after_pt=10); para_keep_with_next(p)
        for para_data in paras:
            num  = para_data.get('number')
            text = para_data.get('text', '')
            if not text.strip(): continue
            blocks = [b.strip() for b in text.split('\n\n') if b.strip()]
            for j, block in enumerate(blocks):
                is_quote = block.startswith('>')
                if is_quote:
                    block = re.sub(r'^>[ \t]*', '', block, flags=re.MULTILINE).strip()
                # Verse detection: multi-line only. Single-line blocks go through
                # add_inline_runs which handles *italic* and **bold** natively.
                is_verse = False
                if '\n' in block:
                    lines_check = [l.strip() for l in block.split('\n') if l.strip()]
                    if is_quote or any(
                        (l.startswith('*') and not l.startswith('**')) or
                        (l.endswith('*') and not l.endswith('**'))
                        for l in lines_check
                    ):
                        is_verse = True
                # Multi-line verse/quote: one para per line, italic, strip per-line *
                if (is_quote or is_verse) and '\n' in block:
                    lines = [l.strip() for l in block.split('\n')]
                    prev_empty = False
                    for line in lines:
                        if not line:
                            prev_empty = True
                            continue
                        line = re.sub(r'^\*(?!\*)', '', line).strip()
                        line = re.sub(r'(?<!\*)\*$', '', line).strip()
                        if not line:
                            prev_empty = True
                            continue
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = p.add_run(line)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(body_pt - 1)
                        run.italic = True
                        para_spacing(p, before_pt=8 if prev_empty else 0, after_pt=2)
                        prev_empty = False
                    continue
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if is_quote else WD_ALIGN_PARAGRAPH.JUSTIFY
                if j == 0 and num is not None:
                    rn = p.add_run(f'{num}.  ')
                    rn.font.name = 'Times New Roman'
                    rn.font.size = Pt(body_pt); rn.bold = True
                add_inline_runs(p, block, base_size_pt=body_pt)
                para_spacing(p, before_pt=0, after_pt=5, line_rule='auto', line_val=276)


# ---------------------------------------------------------------------------
# Pass 2: footnote XML builders
# ---------------------------------------------------------------------------

def _fn_text_runs_xml(text, size_pt, font='Times New Roman'):
    """Return list of lxml w:r elements built from inline-markdown text."""
    runs = []
    last = 0

    def make_r(t, bold=False, italic=False):
        r = etree.Element(_w('r'))
        rPr = etree.SubElement(r, _w('rPr'))
        rf = etree.SubElement(rPr, _w('rFonts'))
        rf.set(_w('ascii'), font); rf.set(_w('hAnsi'), font)
        sz = etree.SubElement(rPr, _w('sz'))
        sz.set(_w('val'), str(int(size_pt * 2)))
        szc = etree.SubElement(rPr, _w('szCs'))
        szc.set(_w('val'), str(int(size_pt * 2)))
        if bold:  etree.SubElement(rPr, _w('b'))
        if italic: etree.SubElement(rPr, _w('i'))
        te = etree.SubElement(r, _w('t'))
        te.set(_XML_SPACE, 'preserve'); te.text = t
        return r

    for m in _INLINE_RE.finditer(text):
        if m.start() > last:
            runs.append(make_r(text[last:m.start()]))
        bi, b, i, fn, fn2 = m.group(2), m.group(3), m.group(4), m.group(5), m.group(6)
        fn_n = fn or fn2
        if bi:   runs.append(make_r(bi, bold=True, italic=True))
        elif b:  runs.append(make_r(b,  bold=True))
        elif i:  runs.append(make_r(i,  italic=True))
        elif fn_n: runs.append(make_r(fn_n))   # nested refs → plain
        last = m.end()
    if last < len(text):
        runs.append(make_r(text[last:]))
    return runs


def _build_footnotes_xml_bytes(fn_dict, note_size_pt=8.5, font='Times New Roman'):
    """Build a complete word/footnotes.xml as UTF-8 bytes."""
    root = etree.Element(_w('footnotes'), nsmap={'w': _WNS, 'r': _RNS})
    # Required separator stubs
    for fn_type, fn_id_str in [('separator', '-1'), ('continuationSeparator', '0')]:
        fe = etree.SubElement(root, _w('footnote'))
        fe.set(_w('type'), fn_type); fe.set(_w('id'), fn_id_str)
        p = etree.SubElement(fe, _w('p'))
        r = etree.SubElement(p, _w('r'))
        etree.SubElement(r, _w(fn_type))
    # Content footnotes
    for num in sorted(fn_dict.keys()):
        fe = etree.SubElement(root, _w('footnote'))
        fe.set(_w('id'), str(num))
        p = etree.SubElement(fe, _w('p'))
        pPr = etree.SubElement(p, _w('pPr'))
        ps = etree.SubElement(pPr, _w('pStyle')); ps.set(_w('val'), 'FootnoteText')
        sp = etree.SubElement(pPr, _w('spacing'))
        sp.set(_w('after'), '0'); sp.set(_w('line'), '240'); sp.set(_w('lineRule'), 'auto')
        # Auto-number mark: 9pt bold inline, no superscript style
        rm = etree.SubElement(p, _w('r'))
        rPr = etree.SubElement(rm, _w('rPr'))
        sz_m = etree.SubElement(rPr, _w('sz')); sz_m.set(_w('val'), '18')   # 9pt
        szc  = etree.SubElement(rPr, _w('szCs')); szc.set(_w('val'), '18')
        etree.SubElement(rPr, _w('b'))
        etree.SubElement(rm, _w('footnoteRef'))
        # Body text
        for re_elem in _fn_text_runs_xml(' ' + fn_dict[num], note_size_pt, font):
            p.append(re_elem)
    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


# ---------------------------------------------------------------------------
# Pass 2: zip manipulation
# ---------------------------------------------------------------------------

def inject_footnotes(docx_path, footnotes, note_size_pt=8.5, font='Times New Roman'):
    """
    In-place: swap superscript-number runs → w:footnoteReference,
    create word/footnotes.xml, patch rels and content types.
    """
    if not footnotes:
        return
    fn_dict = {int(k): v for k, v in footnotes.items()}

    # Read all zip entries
    with zipfile.ZipFile(docx_path, 'r') as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    # ── 1. Patch document.xml ─────────────────────────────────────────────
    doc_tree = etree.fromstring(files['word/document.xml'])
    replaced = 0
    for r_elem in list(doc_tree.findall(f'.//{_w("r")}')):
        rPr = r_elem.find(_w('rPr'))
        if rPr is None: continue
        va = rPr.find(_w('vertAlign'))
        if va is None or va.get(_w('val')) != 'superscript': continue
        t_elem = r_elem.find(_w('t'))
        if t_elem is None or not (t_elem.text or '').strip(): continue
        txt = t_elem.text.strip()
        if not re.match(r'^\d{1,3}$', txt): continue
        fn_id = int(txt)
        if fn_id not in fn_dict: continue
        # Replace run
        new_r = etree.Element(_w('r'))
        new_rPr = etree.SubElement(new_r, _w('rPr'))
        new_rs  = etree.SubElement(new_rPr, _w('rStyle'))
        new_rs.set(_w('val'), 'FootnoteReference')
        new_ref = etree.SubElement(new_r, _w('footnoteReference'))
        new_ref.set(_w('id'), str(fn_id))
        parent = r_elem.getparent()
        idx = list(parent).index(r_elem)
        parent.remove(r_elem)
        parent.insert(idx, new_r)
        replaced += 1
    print(f"[Pass 2] {replaced} superscript runs → footnoteReference")
    files['word/document.xml'] = etree.tostring(
        doc_tree, xml_declaration=True, encoding='UTF-8', standalone=True)

    # ── 2. Build footnotes.xml ────────────────────────────────────────────
    files['word/footnotes.xml'] = _build_footnotes_xml_bytes(fn_dict, note_size_pt, font)
    print(f"[Pass 2] Built footnotes.xml ({len(fn_dict)} footnotes)")

    # ── 3. Patch rels ─────────────────────────────────────────────────────
    rk = 'word/_rels/document.xml.rels'
    rt = etree.fromstring(files[rk])
    if not any(r.get('Type') == _RT_FOOTNOTES
               for r in rt.findall(f'{{{_RELS_NS}}}Relationship')):
        used = {r.get('Id', '') for r in rt.findall(f'{{{_RELS_NS}}}Relationship')}
        rid = 'rIdFN1'
        n = 1
        while rid in used: n += 1; rid = f'rIdFN{n}'
        nr = etree.SubElement(rt, f'{{{_RELS_NS}}}Relationship')
        nr.set('Id', rid); nr.set('Type', _RT_FOOTNOTES); nr.set('Target', 'footnotes.xml')
    files[rk] = etree.tostring(rt, xml_declaration=True, encoding='UTF-8', standalone=True)

    # ── 4. Patch content types ────────────────────────────────────────────
    ct = etree.fromstring(files['[Content_Types].xml'])
    if not any(e.get('PartName') == '/word/footnotes.xml'
               for e in ct.findall(f'{{{_CT_NS}}}Override')):
        nc = etree.SubElement(ct, f'{{{_CT_NS}}}Override')
        nc.set('PartName', '/word/footnotes.xml'); nc.set('ContentType', _CT_FOOTNOTES)
    files['[Content_Types].xml'] = etree.tostring(
        ct, xml_declaration=True, encoding='UTF-8', standalone=True)

    # ── 5. Write patched zip ──────────────────────────────────────────────
    tmp = docx_path + '.tmp_fn'
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)
    os.replace(tmp, docx_path)
    print(f"[Pass 2] ✓ Saved → {docx_path}")


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_booklet_docx(json_path, output_path, size_name='half-letter'):
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    w, h = PAGE_SIZES[size_name]
    m    = MARGINS[size_name]
    doc  = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    section = doc.sections[0]
    set_page_size(section, w, h)
    set_mirror_margins(section, m)
    enable_mirror_margins(doc)

    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].clear()
    section.first_page_footer.paragraphs[0].clear()

    doc_title_short = (data.get('title') or data.get('doc_type', ''))[:45]
    set_header_text(section.header, left_text=doc_title_short, right_text='')
    add_page_number_to_footer(section.footer)

    # Pass 1    build_title_page(doc, data, size_name)
    build_body(doc, data.get('chapters', []), size_name)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"[Pass 1] ✓ Document written ({Path(output_path).stat().st_size:,} bytes)")

    footnotes = data.get('footnotes', {})
    if footnotes:
        inject_footnotes(output_path, footnotes)
    else:
        print("[Pass 2] No footnotes — skipping")

    return output_path


def main():
    p = argparse.ArgumentParser(description='Build booklet .docx from extracted JSON.')
    p.add_argument('json_file')
    p.add_argument('--output', '-o', default=None)
    p.add_argument('--size', '-s', default='half-letter', choices=list(PAGE_SIZES.keys()))
    args = p.parse_args()
    if args.output is None:
        stem = Path(args.json_file).stem.replace('_extracted', '')
        args.output = f"{stem}_booklet.docx"
    print(f"\n[Step 2b] {args.json_file} → {args.output}  (size={args.size})")
    build_booklet_docx(args.json_file, args.output, args.size)
    print(f"[Step 2b] ✓ Done\n")


if __name__ == '__main__':
    main()
